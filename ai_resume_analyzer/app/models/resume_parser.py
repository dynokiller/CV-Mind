"""
app/models/resume_parser.py

Enhanced resume text extractor.

Supported formats:
  - .pdf   : PyMuPDF direct extraction → OCR fallback (scanned PDFs)
  - .docx  : python-docx
  - .jpg / .jpeg / .png / .bmp / .tiff : OCR via ocr_engine

Feature flags (env vars):
  USE_OCR=true (default) — enable OCR for images and scanned PDFs
  USE_OCR=false          — skip OCR; return empty string for image inputs
"""

import os
import logging
from typing import Optional

logger = logging.getLogger("resume_analyzer.parser")

# ── Feature flag ──────────────────────────────────────────────────────────────
USE_OCR = os.getenv("USE_OCR", "true").lower() == "true"


# ═══════════════════════════════════════════════════════════════════════════════
#  Internal parsers
# ═══════════════════════════════════════════════════════════════════════════════

def _extract_text_pdf(pdf_path: str) -> str:
    """
    PDF extraction with OCR fallback.
    1. Try PyMuPDF direct text extraction.
    2. If USE_OCR and result is too short → OCR each page via ocr_engine.
    """
    try:
        import fitz  # PyMuPDF
        doc  = fitz.open(pdf_path)
        text = "".join(page.get_text() for page in doc)
        doc.close()
        text = text.strip()

        if len(text) >= 50:
            logger.info(f"[Parser] PDF direct: {len(text)} chars from '{pdf_path}'.")
            return _clean(text)

        logger.info(f"[Parser] PDF text too short ({len(text)} chars) — trying OCR fallback.")

    except ImportError:
        logger.warning("[Parser] PyMuPDF not installed. Skipping direct PDF extraction.")
        text = ""
    except Exception as e:
        logger.error(f"[Parser] PyMuPDF failed on '{pdf_path}': {e}")
        text = ""

    # OCR fallback
    if USE_OCR:
        try:
            from app.models.ocr_engine import ocr_pdf_pages
            ocr_text = ocr_pdf_pages(pdf_path)
            logger.info(f"[Parser] PDF OCR fallback: {len(ocr_text)} chars.")
            return ocr_text
        except Exception as e:
            logger.error(f"[Parser] PDF OCR fallback failed: {e}")

    return text  # return whatever we have (may be short)


def _extract_text_docx(docx_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        import docx as python_docx
        doc   = python_docx.Document(docx_path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n".join(parts)
        logger.info(f"[Parser] DOCX: {len(text)} chars from '{docx_path}'.")
        return _clean(text)
    except Exception as e:
        logger.error(f"[Parser] DOCX extraction failed: {e}")
        return ""


def _extract_text_image(image_path: str) -> str:
    """Extract text from an image file using the OCR engine."""
    if not USE_OCR:
        logger.warning(f"[Parser] USE_OCR=false — skipping image OCR for '{image_path}'.")
        return ""
    try:
        from app.models.ocr_engine import extract_text_ocr
        text = extract_text_ocr(image_path)
        logger.info(f"[Parser] Image OCR: {len(text)} chars from '{image_path}'.")
        return text
    except Exception as e:
        logger.error(f"[Parser] Image OCR failed on '{image_path}': {e}")
        return ""


def _clean(text: str) -> str:
    """Light cleanup applied after direct text extraction (not needed for OCR output)."""
    import re
    # Collapse repeated blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Normalise whitespace per line
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(lines).strip()


# ═══════════════════════════════════════════════════════════════════════════════
#  Public API
# ═══════════════════════════════════════════════════════════════════════════════

def parse_resume(file_path: str) -> str:
    """
    Extract plain text from a resume file.

    Supports: .pdf, .docx, .doc, .jpg, .jpeg, .png, .bmp, .tiff

    Args:
        file_path: Absolute or relative path to the resume file.

    Returns:
        Extracted text string. Returns "" if extraction fails.
    """
    if not os.path.exists(file_path):
        logger.error(f"[Parser] File not found: '{file_path}'")
        return ""

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_text_pdf(file_path)

    elif ext in (".docx", ".doc"):
        if ext == ".doc":
            logger.warning(f"[Parser] Legacy .doc file detected: '{file_path}'. This may fail; .docx is preferred.")
        return _extract_text_docx(file_path)

    elif ext in (".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".webp"):
        return _extract_text_image(file_path)

    else:
        logger.error(f"[Parser] Unsupported file extension: '{ext}' for file '{file_path}'")
        return ""


# Backward-compatible alias (used by old imports)
extract_text = parse_resume
